import json
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn


def create_docx_file(data):

    # Create a new Word document
    doc = Document("./template.docx")
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(15)
    section.right_margin = Mm(15)
    section.top_margin = Mm(15.4)
    section.bottom_margin = Mm(0)
    section.footer_distance = Mm(0)

    # Create a table with two columns
    table0 = doc.add_table(rows=1, cols=2)

    # Define the table cells
    cell_00 = table0.cell(0, 0)
    cell_01 = table0.cell(0, 1)

    # Add logo
    image_path = "SMD_icon.PNG"
    cell_00.add_paragraph().add_run().add_picture(image_path)

    # Add name
    name = cell_01.add_paragraph(data['candidate_name'])
    name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    name_runs = name.runs
    for run in name_runs:
        run.font.size = Pt(24)
        run.font.color.theme_color = WD_COLOR_INDEX.DARK_BLUE
        run.font.bold = True

    # Add title
    title = cell_01.add_paragraph(data['candidate_title'])
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_runs = title.runs
    for run in title_runs:
        run.font.size = Pt(16)
        run.font.bold = True

    # Add summary
    cell_01.add_paragraph(data['summary'])

    # Add a horizontal line
    doc.add_paragraph('_' * 90).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph('www.smartdev.com | Hanoi City, Vietnam').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph('_' * 90).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Iterate through all paragraphs and remove space after each one
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.space_after = 0 

    # Create a table with two columns
    table1 = doc.add_table(rows=1, cols=2)

    # Define the table cells
    cell_10 = table1.cell(0, 0)
    cell_11 = table1.cell(0, 1)
    
    cell_10.width = Mm(70)
    cell_11.width = Mm(110)

    # Add links
    if len(data['links']) > 1:
        contact = cell_10.add_paragraph('Contact')
        contact_runs = contact.runs
        for run in contact_runs:
            run.font.size = Pt(18)
            run.font.bold = True

        for link in data['links']:
            cell_10.add_paragraph(link, style='Huy List')

    # Add work experience
    exp = cell_11.add_paragraph('Work Experience')
    exp_runs = exp.runs
    for run in exp_runs:
        run.font.size = Pt(18)
        run.font.bold = True

    for exp in data['work_exp']:
        details = cell_11.add_paragraph(f"{exp['work_timeline'][0]} - {exp['work_timeline'][1]} | {exp['work_title']} - {exp['work_company']}")
        details_runs = details.runs
        for run in details_runs:
            run.font.bold = True
            run.font.size = Pt(14)

        cell_11.add_paragraph(exp['work_description'])

        # Add responsibilities
        res = cell_11.add_paragraph('Responsibilities: ')
        res_runs = res.runs
        for run in res_runs:
            run.font.bold = True
            run.font.size = Pt(12)

        for resp in exp['work_responsibilities']:
            cell_11.add_paragraph(resp, style='Huy List')
        
        # Add technologies
        if len(exp['work_technologies']) > 0:
            techs = cell_11.add_paragraph('Technologies: ')
            techs_runs = techs.runs
            for run in techs_runs:
                run.font.bold = True
                run.font.size = Pt(12)

            cell_11.add_paragraph(exp['work_technologies'])

    # Add projects
    if len(data['projects']) > 0:
        projects = cell_11.add_paragraph('Projects')
        projects_runs = projects.runs
        for run in projects_runs:
            run.font.size = Pt(18)
            run.font.bold = True

        for project in data['projects']:
            details = cell_11.add_paragraph(f"{project['project_timeline'][0]} - {project['project_timeline'][1]} | {project['project_name']}")
            details_runs = details.runs
            for run in details_runs:
                run.font.bold = True
                run.font.size = Pt(14)

            cell_11.add_paragraph(project['project_description'])

            # Add responsibilities
            res = cell_11.add_paragraph('Responsibilities: ')
            res_runs = res.runs
            for run in res_runs:
                run.font.bold = True
                run.font.size = Pt(12)

            for resp in project['project_responsibilities']:
                cell_11.add_paragraph(resp, style='Huy List')
            
            # Add technologies
            if len(project['project_technologies']) > 0:
                techs = cell_11.add_paragraph('Technologies: ')
                techs_runs = techs.runs
                for run in techs_runs:
                    run.font.bold = True
                    run.font.size = Pt(12)

                cell_11.add_paragraph(project['project_technologies'])

    # Add education
    ed = cell_10.add_paragraph('Education')
    edu_runs = ed.runs
    for run in edu_runs:
        run.font.size = Pt(18)
        run.font.bold = True

    for edu in data['education']:
        edu_details = cell_10.add_paragraph(f"{edu['edu_timeline'][0]} - {edu['edu_timeline'][1]} | {edu['edu_degree']} - {edu['edu_school']}")
        edu_details_runs = edu_details.runs
        for run in edu_details_runs:
            run.font.bold = True
    
    # Add languages
    lang = cell_10.add_paragraph('Languages')
    lang_runs = lang.runs
    for run in lang_runs:
        run.font.size = Pt(18)
        run.font.bold = True

    for lang in data['languages']:
        cell_10.add_paragraph(f"{lang['lang']}: {lang['lang_lvl']}", style='Huy List')

    # Add certifications
    if len(data['certifications']) > 0:
        certs = cell_10.add_paragraph('Certifications')
        certs_runs = certs.runs
        for run in certs_runs:
            run.font.size = Pt(18)
            run.font.bold = True

        for cert in data['certifications']:
            cell_10.add_paragraph(cert, style='Huy List')

    # Add skills
    if len(data['skills']) > 0:
        skills = cell_10.add_paragraph('Skills')
        skills_runs = skills.runs
        for run in skills_runs:
            run.font.size = Pt(18)
            run.font.bold = True

        for skill in data['skills']:
            skill_name = skill['skill_name']
            yoe = skill['yoe']
            if yoe is not None:
                skill_text = f"{skill_name} - {yoe} years"
            else:
                skill_text = skill_name
            cell_10.add_paragraph(skill_text, style='Huy List')


    # Add inner vertical border
    borders = OxmlElement('w:tblBorders')
    bottom_border = OxmlElement('w:insideV')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '20')
    borders.append(bottom_border)
    table1._tbl.tblPr.append(borders)
    return doc

def remove_none_value(d):
    if isinstance(d, dict):
        for k,v in d.items():
            if v is None:
                d[k] = ""
            elif isinstance(v, (dict, list)]):
                d[k] = remove_none_value(v)

    elif isinstance(d, list):
        for i,c in enumerate(d):
            if c is None:
                d[i] = ""
            elif isinstance(v, [dict, list]):
                d[i] = remove_none_value(c) 
    return d

def post_process(data):
    # skills = {}

    # for skill in data['skills']:
    #     skill_name = skill['skill_name']
    #     yoe = skill['yoe']

    #     if skill_name in skills:
    #         if yoe is not None and skills[skill_name]['yoe'] is None:
    #             skills[skill_name] = skill
    #         elif yoe is not None and skills[skill_name]['yoe'] is not None and yoe > skills[skill_name]['yoe']:
    #             skills[skill_name] = skill
    #     else:
    #         skills[skill_name] = skill

    # # Convert back to a list
    # final_skills = list(skills.values())

    # # Update skills field in data
    # data['skills'] = final_skills

    """
    remove NONE
    """
    data = remove_none_value(data)

    return data

if __name__ == '__main__':
    with open('.\export_doc.json', 'r') as f:
        data = json.load(f)
    data = post_process(data)
    doc = create_docx_file(data)
    doc.save('export_resume.docx')
